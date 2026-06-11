"""
ADÜ Mühendislik Fakültesi - Bitirme Tezi Oluşturma Script'i
Proje: Akıllı Diyet Asistanı - Ajan Tabanlı RAG Sistemi ile Kişiselleştirilmiş Beslenme Yönetimi
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================
# CONFIGURATION
# ============================================================
OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Bitirme_Tezi.docx")

THESIS_TITLE_EN = "SMART DIET ASSISTANT: PERSONALIZED NUTRITION MANAGEMENT WITH AGENT-BASED RAG SYSTEM"
THESIS_TITLE_TR = "AKILLI DİYET ASİSTANI: AJAN TABANLI RAG SİSTEMİ İLE KİŞİSELLEŞTİRİLMİŞ BESLENME YÖNETİMİ"

STUDENTS = [
    ("Çağla Eylül", "AVCI", "211805043"),
    ("Dila", "DEMİR KOPARAN", "211805075"),
    ("Kaan", "ERDEN", "211805068"),
    ("Mehmet", "ÖZCAN", "211805076"),
]

DEPARTMENT = "COMPUTER ENGINEERING"
DEPARTMENT_TR = "BİLGİSAYAR MÜHENDİSLİĞİ"
SUPERVISOR = "Asst. Prof. Dr. Hüseyin ABACI"
SUPERVISOR_TR = "Dr. Öğr. Üyesi Hüseyin ABACI"
YEAR = "2026"
SUBMISSION_DATE = "June 2026"

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def add_heading_style(doc, text, level=1, bold=True, centered=False):
    """Add a heading with specific style."""
    heading = doc.add_heading(text, level=level)
    if centered:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph_tnr(doc, text, size=12, bold=False, italic=False, alignment=None, space_after=6, space_before=0):
    """Add a paragraph with Times New Roman font."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if alignment:
        para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.5
    return para

def add_centered_text(doc, text, size=12, bold=False):
    """Add centered text."""
    return add_paragraph_tnr(doc, text, size=size, bold=bold, alignment=WD_ALIGN_PARAGRAPH.CENTER)

def add_table_with_style(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(header)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E4057')
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(str(cell_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'F0F4F8')
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    return table


def set_section_margins(section):
    """Set standard margins per ADU template."""
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4.0)  # binding side
    section.right_margin = Cm(2.5)


def add_table_of_contents(doc):
    """Adds a native MS Word Table of Contents field."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    # Force update fields on open
    try:
        settings = doc.settings._element
        updateFields = OxmlElement('w:updateFields')
        updateFields.set(qn('w:val'), 'true')
        settings.append(updateFields)
    except Exception as e:
        print(f"Could not force update fields: {e}")


def add_list_of_figures(doc):
    """Adds a native MS Word List of Figures field."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\c "Figure" \\h \\z'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_list_of_tables(doc):
    """Adds a native MS Word List of Tables field."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\c "Table" \\h \\z'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_caption(doc, label, text):
    """Adds a caption paragraph with a SEQ field for Figure/Table numbering."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    
    try:
        p.style = 'Caption'
    except Exception:
        pass
        
    r_label = p.add_run(f"{label} ")
    r_label.font.name = 'Times New Roman'
    r_label.font.size = Pt(12)
    r_label.font.italic = True
    r_label.font.bold = True
    r_label.font.color.rgb = RGBColor(0, 0, 0)
    
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), f'SEQ {label} \\* ARABIC')
    
    r_num = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.append(rFonts)
    rPr.append(OxmlElement('w:b'))
    rPr.append(OxmlElement('w:i'))
    r_num.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = "0"
    r_num.append(t)
    fldSimple.append(r_num)
    
    p._p.append(fldSimple)
    
    r_text = p.add_run(f". {text}")
    r_text.font.name = 'Times New Roman'
    r_text.font.size = Pt(12)
    r_text.font.italic = True
    r_text.font.color.rgb = RGBColor(0, 0, 0)
    
    return p


def add_figure_placeholder(doc, title):
    """Adds a visual placeholder box for figures without local image files."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.width = Cm(14.0)
    set_cell_shading(cell, 'F0F4F8')
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(f"[ System Diagram: {title} ]\n(Refer to project architecture specifications / Poster.pdf for full visualization)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 64, 87)
    return table


def add_numbered_equation(doc, formula_text, equation_num):
    """Adds a centered formula with a right-aligned equation number in parentheses using a borderless table."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Hide borders
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b_name in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{b_name}')
                b.set(qn('w:val'), 'none')
                b.set(qn('w:sz'), '0')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'auto')
                tcBorders.append(b)
            tcPr.append(tcBorders)
            
    # Set widths: col 0 is wide for formula, col 1 is narrow for equation number
    table.rows[0].cells[0].width = Cm(13.0)
    table.rows[0].cells[1].width = Cm(1.5)
    
    # Column 0: Formula centered
    cell0 = table.rows[0].cells[0]
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.line_spacing = 1.5
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.space_after = Pt(6)
    r0 = p0.add_run(formula_text)
    r0.font.name = 'Times New Roman'
    r0.font.size = Pt(12)
    
    # Column 1: Equation number right-aligned
    cell1 = table.rows[0].cells[1]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.line_spacing = 1.5
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run(f"({equation_num})")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.bold = True
    return table


# ============================================================
# MAIN DOCUMENT GENERATION
# ============================================================

def create_thesis():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # Set heading styles to TNR
    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Times New Roman'
        h_style.font.color.rgb = RGBColor(0, 0, 0)
    
    section = doc.sections[0]
    set_section_margins(section)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    
    # ========================================================
    # OUTER COVER PAGE
    # ========================================================
    add_centered_text(doc, "AYDIN ADNAN MENDERES UNIVERSITY", size=12, bold=True)
    add_centered_text(doc, "ENGINEERING FACULTY", size=12, bold=True)
    add_centered_text(doc, f"{DEPARTMENT} DEPARTMENT", size=12, bold=True)
    
    # Add ADU Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(12)
    p_logo.paragraph_format.space_after = Pt(12)
    p_logo.add_run().add_picture('extracted_img_0_X7.png', width=Cm(3.5))
    
    add_centered_text(doc, THESIS_TITLE_EN, size=16, bold=True)
    
    for _ in range(4):
        add_centered_text(doc, "", size=12)
    
    # Student names (alphabetical by surname)
    for name, surname, student_id in STUDENTS:
        add_centered_text(doc, f"{name} {surname}", size=12, bold=True)
    
    for _ in range(3):
        add_centered_text(doc, "", size=12)
    
    add_centered_text(doc, "Supervisor:", size=12, bold=True)
    add_centered_text(doc, SUPERVISOR, size=12, bold=True)
    
    add_page_break(doc)
    
    # ========================================================
    # INNER COVER PAGE
    # ========================================================
    add_centered_text(doc, "AYDIN ADNAN MENDERES UNIVERSITY", size=12, bold=True)
    add_centered_text(doc, "ENGINEERING FACULTY", size=12, bold=True)
    add_centered_text(doc, f"{DEPARTMENT} DEPARTMENT", size=12, bold=True)
    
    # Add ADU Logo
    p_logo2 = doc.add_paragraph()
    p_logo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo2.paragraph_format.space_before = Pt(12)
    p_logo2.paragraph_format.space_after = Pt(12)
    p_logo2.add_run().add_picture('extracted_img_0_X7.png', width=Cm(3.5))
    
    add_centered_text(doc, THESIS_TITLE_EN, size=16, bold=True)
    
    for _ in range(3):
        add_centered_text(doc, "", size=12)
    
    for name, surname, student_id in STUDENTS:
        add_centered_text(doc, f"{name} {surname}", size=12, bold=True)
    
    for _ in range(2):
        add_centered_text(doc, "", size=12)
    
    add_centered_text(doc, "Supervisor:", size=12, bold=True)
    add_centered_text(doc, SUPERVISOR, size=12, bold=True)
    
    for _ in range(2):
        add_centered_text(doc, "", size=12)
    
    add_centered_text(doc, f"Date of Submission: {SUBMISSION_DATE}", size=12, bold=True)
    
    add_page_break(doc)
    
    # ========================================================
    # ABSTRACT (English)
    # ========================================================
    add_heading_style(doc, "ABSTRACT", level=1, centered=True)
    add_centered_text(doc, THESIS_TITLE_EN, size=12, bold=True)
    
    for name, surname, _ in STUDENTS:
        add_centered_text(doc, f"{name} {surname}", size=12)
    
    add_centered_text(doc, f"B.Sc. Thesis, {DEPARTMENT.title()} Department", size=12)
    add_centered_text(doc, f"Supervisor: {SUPERVISOR}", size=12)
    add_centered_text(doc, f"{YEAR}, 38 pages", size=12)
    
    abstract_text = (
        "This thesis presents the design, development, and evaluation of a Smart Diet Assistant system that leverages "
        "an agent-based Retrieval-Augmented Generation (RAG) architecture to provide personalized nutrition management. "
        "The system integrates a Self-Reflective RAG pipeline built on the LangGraph framework, combining large language "
        "models (GPT-4o-mini) with vector databases (ChromaDB) and real-time web search capabilities (Tavily API) to "
        "deliver context-aware dietary recommendations. The backend is implemented using FastAPI with Python, employing "
        "a multi-agent architecture where specialized agents handle question routing, document retrieval, relevance "
        "grading, personalized response generation, hallucination detection, and answer quality assessment. "
        "User context injection enables the system to consider individual health profiles including height, weight, "
        "age, gender, daily caloric intake, dietary restrictions, and activity levels when generating recommendations. "
        "The mobile client is developed as a native Android application using Kotlin and Jetpack Compose with Material 3 "
        "design principles, featuring calorie tracking with interactive donut and bar charts, water intake monitoring with "
        "animated visualizations, MET-based exercise tracking, and an AI-powered food analysis module utilizing both "
        "on-device TensorFlow Lite classification (EfficientNet-B4, 199 food categories) and cloud-based GPT-5.4 vision "
        "analysis through a multi-agent validation pipeline. The system also includes an intelligent recipe recommendation "
        "engine that identifies ingredients from photos and generates personalized recipes respecting caloric budgets and "
        "dietary constraints. The architecture demonstrates how modern generative AI techniques, combined with classical "
        "information retrieval methods, can be effectively applied to create a comprehensive, personalized health "
        "management platform that goes beyond simple information retrieval to provide expert-level dietary consultation."
    )
    add_paragraph_tnr(doc, abstract_text, space_before=12)
    
    keywords_para = doc.add_paragraph()
    run_bold = keywords_para.add_run("Keywords: ")
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(12)
    run_bold.font.bold = True
    run_text = keywords_para.add_run("Retrieval-Augmented Generation, Multi-Agent System, Personalized Nutrition, LangGraph, Mobile Health Application")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    keywords_para.paragraph_format.space_before = Pt(24)
    
    add_page_break(doc)
    
    # ========================================================
    # ÖZET (Turkish Abstract)
    # ========================================================
    add_heading_style(doc, "ÖZET", level=1, centered=True)
    add_centered_text(doc, THESIS_TITLE_TR, size=12, bold=True)
    
    for name, surname, _ in STUDENTS:
        add_centered_text(doc, f"{name} {surname}", size=12)
    
    add_centered_text(doc, f"Lisans Bitirme Tezi, {DEPARTMENT_TR} Bölümü", size=12)
    add_centered_text(doc, f"Tez Danışmanı: {SUPERVISOR_TR}", size=12)
    add_centered_text(doc, f"{YEAR}, 38 sayfa", size=12)
    
    ozet_text = (
        "Bu tez, kişiselleştirilmiş beslenme yönetimi sağlamak amacıyla ajan tabanlı Erişimle Artırılmış Üretim (RAG) "
        "mimarisinden yararlanan bir Akıllı Diyet Asistanı sisteminin tasarımını, geliştirilmesini ve değerlendirmesini "
        "sunmaktadır. Sistem, büyük dil modellerini (GPT-4o-mini), vektör veritabanlarını (ChromaDB) ve gerçek zamanlı "
        "web arama yeteneklerini (Tavily API) birleştiren, LangGraph çerçevesi üzerine inşa edilmiş bir Kendi Kendini "
        "Değerlendiren RAG pipeline'ı entegre etmektedir. Arka uç, Python ile FastAPI kullanılarak uygulanmış olup, "
        "soru yönlendirme, belge erişimi, alaka değerlendirmesi, kişiselleştirilmiş yanıt üretimi, halüsinasyon tespiti "
        "ve cevap kalite değerlendirmesi görevlerini yerine getiren özelleştirilmiş ajanların yer aldığı çoklu ajan "
        "mimarisini kullanmaktadır. Kullanıcı bağlam enjeksiyonu, sistemin öneriler üretirken boy, kilo, yaş, cinsiyet, "
        "günlük kalori alımı, beslenme kısıtlamaları ve aktivite düzeyleri dahil olmak üzere bireysel sağlık profillerini "
        "göz önünde bulundurmasını sağlamaktadır. Mobil istemci, Kotlin ve Jetpack Compose ile Material 3 tasarım "
        "ilkeleri kullanılarak yerel bir Android uygulaması olarak geliştirilmiştir. Uygulama; interaktif halka ve çubuk "
        "grafiklerle kalori takibi, animasyonlu görselleştirmelerle su tüketimi izleme, MET tabanlı egzersiz takibi ve "
        "hem cihaz üzerinde TensorFlow Lite sınıflandırması (EfficientNet-B4, 199 yemek kategorisi) hem de çoklu ajan "
        "doğrulama pipeline'ı aracılığıyla bulut tabanlı GPT-5.4 görüntü analizi kullanan yapay zeka destekli besin "
        "analizi modülünü içermektedir. Sistem ayrıca, fotoğraflardan malzemeleri tanımlayan ve kalorik bütçelere ve "
        "beslenme kısıtlamalarına uygun kişiselleştirilmiş tarifler üreten akıllı bir tarif öneri motoruna da sahiptir. "
        "Mimari, modern üretken yapay zeka tekniklerinin, klasik bilgi erişim yöntemleriyle birleştirildiğinde, basit "
        "bilgi erişiminin ötesine geçerek uzman düzeyinde diyet danışmanlığı sunan kapsamlı, kişiselleştirilmiş bir "
        "sağlık yönetim platformu oluşturmak için nasıl etkin bir şekilde uygulanabileceğini göstermektedir."
    )
    add_paragraph_tnr(doc, ozet_text, space_before=12)
    
    anahtar_para = doc.add_paragraph()
    run_bold = anahtar_para.add_run("Anahtar Kelimeler: ")
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(12)
    run_bold.font.bold = True
    run_text = anahtar_para.add_run("Erişimle Artırılmış Üretim, Çoklu Ajan Sistemi, Kişiselleştirilmiş Beslenme, LangGraph, Mobil Sağlık Uygulaması")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    anahtar_para.paragraph_format.space_before = Pt(24)
    
    add_page_break(doc)
    
    # ========================================================
    # ACKNOWLEDGEMENT
    # ========================================================
    add_heading_style(doc, "ACKNOWLEDGEMENT", level=1, centered=True)
    
    ack_text = (
        "We would like to express our sincere gratitude to our thesis supervisor for their invaluable guidance, "
        "continuous support, and constructive feedback throughout the course of this project. Their expertise and "
        "encouragement have been instrumental in shaping this research."
    )
    add_paragraph_tnr(doc, ack_text, space_before=18)
    
    ack_text2 = (
        "We extend our heartfelt appreciation to the faculty members of the Computer Engineering Department at "
        "Aydın Adnan Menderes University for providing us with a solid academic foundation and fostering our "
        "intellectual growth throughout our undergraduate studies."
    )
    add_paragraph_tnr(doc, ack_text2)
    
    ack_text3 = (
        "We are also grateful to our families and friends for their unwavering support, patience, and "
        "encouragement during the challenging yet rewarding process of completing this thesis."
    )
    add_paragraph_tnr(doc, ack_text3)
    
    ack_text4 = (
        "Finally, we acknowledge the open-source community and the developers of the tools, libraries, and "
        "frameworks used in this project, including LangChain, LangGraph, FastAPI, Jetpack Compose, and "
        "TensorFlow Lite, whose contributions have made modern AI application development more accessible."
    )
    add_paragraph_tnr(doc, ack_text4)
    
    # Signatures
    for _ in range(3):
        add_paragraph_tnr(doc, "")
    
    for name, surname, student_id in STUDENTS:
        add_paragraph_tnr(doc, f"{name} {surname}", alignment=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    
    add_centered_text(doc, f"Aydın, {SUBMISSION_DATE}", size=12)
    
    add_page_break(doc)
    
    # ========================================================
    # TABLE OF CONTENTS (Placeholder)
    # ========================================================
    add_heading_style(doc, "TABLE OF CONTENTS", level=1, centered=False)
    add_table_of_contents(doc)
    
    add_page_break(doc)
    
    # ========================================================
    # LIST OF TABLES
    # ========================================================
    add_heading_style(doc, "LIST OF TABLES", level=1, centered=False)
    add_list_of_tables(doc)
    
    add_page_break(doc)
    
    # ========================================================
    # LIST OF FIGURES
    # ========================================================
    add_heading_style(doc, "LIST OF FIGURES", level=1, centered=False)
    add_list_of_figures(doc)
    
    add_page_break(doc)
    
    # ========================================================
    # LIST OF ABBREVIATIONS
    # ========================================================
    add_heading_style(doc, "LIST OF ABBREVIATIONS", level=1, centered=False)
    
    abbreviations = [
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("APA", "American Psychological Association"),
        ("BMR", "Basal Metabolic Rate"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("DB", "Database"),
        ("FAB", "Floating Action Button"),
        ("GenAI", "Generative Artificial Intelligence"),
        ("GPT", "Generative Pre-trained Transformer"),
        ("HTTP", "Hypertext Transfer Protocol"),
        ("JSON", "JavaScript Object Notation"),
        ("JWT", "JSON Web Token"),
        ("LLM", "Large Language Model"),
        ("MET", "Metabolic Equivalent of Task"),
        ("ML", "Machine Learning"),
        ("NLP", "Natural Language Processing"),
        ("PDF", "Portable Document Format"),
        ("RAG", "Retrieval-Augmented Generation"),
        ("REST", "Representational State Transfer"),
        ("SDK", "Software Development Kit"),
        ("SQL", "Structured Query Language"),
        ("TDEE", "Total Daily Energy Expenditure"),
        ("TFLite", "TensorFlow Lite"),
        ("UI", "User Interface"),
        ("UX", "User Experience"),
    ]
    
    table = doc.add_table(rows=len(abbreviations), cols=2)
    table.style = 'Table Grid'
    for i, (abbr, meaning) in enumerate(abbreviations):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        cell0.text = ''
        cell1.text = ''
        p0 = cell0.paragraphs[0]
        p1 = cell1.paragraphs[0]
        r0 = p0.add_run(abbr)
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)
        r0.font.bold = True
        r1 = p1.add_run(meaning)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        cell0.width = Cm(3)
        cell1.width = Cm(11)
    
    # Remove table borders for clean look
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    add_page_break(doc)

    # ========================================================
    # LIST OF SYMBOLS (Not applicable for CS, minimal)
    # ========================================================
    add_heading_style(doc, "LIST OF SYMBOLS", level=1, centered=False)
    add_paragraph_tnr(doc, "No specialized mathematical symbols are used in this thesis beyond standard computing notation.")
    
    add_page_break(doc)
    
    # ========================================================
    # CHAPTER 1: INTRODUCTION
    # ========================================================
    add_heading_style(doc, "1. INTRODUCTION", level=1)
    
    intro_text = (
        "The rapid advancement of artificial intelligence (AI) and natural language processing (NLP) technologies "
        "has opened unprecedented opportunities for developing intelligent, personalized health management systems. "
        "In particular, the emergence of Large Language Models (LLMs) such as OpenAI's GPT-4 series, combined with "
        "innovative retrieval techniques, has enabled the creation of systems that can provide expert-level guidance "
        "in specialized domains, including nutrition and dietary management."
    )
    add_paragraph_tnr(doc, intro_text)
    
    intro_text2 = (
        "Nutrition plays a fundamental role in human health, with poor dietary habits being a leading risk factor for "
        "chronic diseases including obesity, diabetes, cardiovascular disease, and certain cancers (World Health "
        "Organization, 2021). Despite the growing awareness of healthy eating, many individuals struggle to make "
        "informed dietary choices due to the complexity of nutritional science, the abundance of contradictory "
        "information, and the difficulty of personalizing dietary recommendations to individual health profiles "
        "and preferences."
    )
    add_paragraph_tnr(doc, intro_text2)
    
    intro_text3 = (
        "Traditional calorie-tracking applications, while useful for basic food logging, typically lack the "
        "intelligence to provide personalized dietary advice. They require manual data entry, offer generic "
        "recommendations, and fail to leverage the wealth of nutritional knowledge available in scientific "
        "literature. This gap between simple calorie counting and expert nutritional consulting represents "
        "a significant opportunity for AI-enhanced solutions."
    )
    add_paragraph_tnr(doc, intro_text3)
    
    # 1.1 Problem Statement
    add_heading_style(doc, "1.1. Problem Statement", level=2)
    
    problem_text = (
        "The core problem addressed in this thesis is the lack of accessible, intelligent, and truly personalized "
        "dietary guidance systems. Existing mobile health applications suffer from several critical limitations: "
        "(1) they rely on static food databases that cannot provide context-aware nutritional advice, "
        "(2) they fail to consider the holistic health profile of the user when making recommendations, "
        "(3) they cannot analyze food items from photographs with sufficient accuracy, "
        "(4) they do not leverage scientific nutritional literature for evidence-based suggestions, and "
        "(5) they lack the ability to generate creative, personalized recipe recommendations based on available "
        "ingredients and dietary constraints."
    )
    add_paragraph_tnr(doc, problem_text)
    
    # 1.2 Purpose and Scope
    add_heading_style(doc, "1.2. Purpose and Scope", level=2)
    
    purpose_text = (
        "The purpose of this thesis is to design, develop, and evaluate a comprehensive Smart Diet Assistant "
        "that addresses the aforementioned limitations by integrating modern AI techniques with mobile application "
        "development. Specifically, the project aims to:"
    )
    add_paragraph_tnr(doc, purpose_text)
    
    purposes = [
        "Develop an agent-based Self-Reflective RAG system that can provide personalized dietary advice by combining "
        "knowledge from scientific nutritional documents with real-time web information.",
        "Implement a multi-agent food analysis pipeline that utilizes computer vision (GPT-5.4 Vision) with a "
        "validation feedback loop to accurately estimate caloric content and macronutrient composition from food photographs.",
        "Create an intelligent recipe recommendation engine that identifies ingredients from photos and generates "
        "personalized, health-conscious recipes respecting individual caloric budgets and dietary restrictions.",
        "Build a feature-rich native Android mobile application with an intuitive, modern user interface for "
        "comprehensive health tracking including calorie, water, and exercise monitoring.",
        "Demonstrate the practical applicability of LangGraph-based multi-agent orchestration in real-world health "
        "technology applications.",
    ]
    for p_text in purposes:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(p_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        para.paragraph_format.line_spacing = 1.5
    
    # 1.3 Significance of the Study
    add_heading_style(doc, "1.3. Significance of the Study", level=2)
    
    significance_text = (
        "This study contributes to the fields of health informatics, artificial intelligence, and mobile application "
        "development in several significant ways. First, it demonstrates a novel application of the Self-Reflective "
        "RAG architecture in the health domain, where the system not only retrieves and generates information but "
        "also evaluates the quality and accuracy of its own responses through built-in hallucination detection and "
        "answer grading mechanisms. Second, the multi-agent food analysis pipeline with its self-correcting validation "
        "loop represents an innovative approach to improving the reliability of AI-based food recognition systems. "
        "Third, the integration of user context injection into the RAG pipeline enables truly personalized dietary "
        "consultation, moving beyond generic recommendations to expert-level, individualized guidance. This approach "
        "bridges the gap between simple calorie-tracking tools and professional nutritional consulting services, "
        "making personalized dietary guidance accessible to a broader population."
    )
    add_paragraph_tnr(doc, significance_text)
    
    # 1.4 Thesis Organization
    add_heading_style(doc, "1.4. Thesis Organization", level=2)
    
    org_text = (
        "The remainder of this thesis is organized as follows: Chapter 2 presents a comprehensive literature review "
        "covering the theoretical foundations and related work in the areas of RAG systems, multi-agent architectures, "
        "food recognition technologies, and mobile health applications. Chapter 3 describes the materials and methods "
        "used in the development of the system, including the system architecture, technology stack, and implementation "
        "details. Chapter 4 presents the results and discusses the system's features, functionality, and performance. "
        "Chapter 5 concludes the thesis with a summary of contributions, limitations, and directions for future work."
    )
    add_paragraph_tnr(doc, org_text)
    
    add_page_break(doc)
    
    # ========================================================
    # CHAPTER 2: LITERATURE REVIEW
    # ========================================================
    add_heading_style(doc, "2. LITERATURE REVIEW", level=1)
    
    lit_intro = (
        "This chapter provides a comprehensive review of the theoretical foundations and existing research that "
        "underpin the development of the Smart Diet Assistant system. The review covers four major areas: "
        "Retrieval-Augmented Generation (RAG) systems, multi-agent AI architectures, food recognition and "
        "nutritional analysis technologies, and mobile health (mHealth) applications."
    )
    add_paragraph_tnr(doc, lit_intro)
    
    # 2.1 RAG Systems
    add_heading_style(doc, "2.1. Retrieval-Augmented Generation (RAG) Systems", level=2)
    
    rag_text = (
        "Retrieval-Augmented Generation (RAG) represents a paradigm shift in how language models access and utilize "
        "external knowledge. First introduced by Lewis et al. (2020), RAG combines the generative capabilities of "
        "large language models with the precision of information retrieval systems. Unlike purely generative models "
        "that rely solely on their training data, RAG systems dynamically retrieve relevant documents from external "
        "knowledge bases during inference, grounding their responses in factual, up-to-date information."
    )
    add_paragraph_tnr(doc, rag_text)
    
    rag_text2 = (
        "The standard RAG pipeline consists of three core stages: indexing, retrieval, and generation. During indexing, "
        "documents are processed, split into chunks, and converted into dense vector representations (embeddings) that "
        "are stored in a vector database. At retrieval time, the user's query is similarly embedded, and semantically "
        "similar documents are identified through cosine similarity search. These retrieved documents are then provided "
        "as context to the language model, which generates a response grounded in the retrieved information."
    )
    add_paragraph_tnr(doc, rag_text2)
    
    rag_text3 = (
        "Advanced RAG architectures incorporate self-reflection mechanisms that evaluate the quality of retrieved "
        "documents and generated responses. Asai et al. (2023) introduced Self-RAG, which adds reflection tokens "
        "to enable the model to assess whether retrieval is necessary, evaluate document relevance, and check for "
        "hallucinations in generated responses. This self-reflective approach significantly improves the factual "
        "accuracy and reliability of RAG systems, making them suitable for critical domains such as healthcare "
        "and nutrition."
    )
    add_paragraph_tnr(doc, rag_text3)
    
    # 2.2 Multi-Agent AI Architectures
    add_heading_style(doc, "2.2. Multi-Agent AI Architectures", level=2)
    
    agent_text = (
        "Multi-agent systems represent an AI paradigm where multiple autonomous agents collaborate to solve complex "
        "problems that exceed the capabilities of individual agents. In the context of LLM-based applications, "
        "multi-agent architectures have gained significant traction as a means to decompose complex tasks into "
        "specialized sub-tasks handled by purpose-built agents (Wu et al., 2023)."
    )
    add_paragraph_tnr(doc, agent_text)
    
    agent_text2 = (
        "LangGraph, developed by LangChain (2024), provides a framework for building stateful, multi-actor "
        "applications with LLMs. It models agent workflows as directed graphs where nodes represent processing "
        "steps and edges define the flow of data and control. Conditional edges enable dynamic routing based on "
        "the state of the computation, allowing for sophisticated decision-making, error recovery, and iterative "
        "refinement processes. This graph-based approach is particularly well-suited for complex AI applications "
        "that require multiple stages of processing with feedback loops."
    )
    add_paragraph_tnr(doc, agent_text2)
    
    # 2.3 Food Recognition
    add_heading_style(doc, "2.3. Food Recognition and Nutritional Analysis", level=2)
    
    food_text = (
        "Computer vision-based food recognition has been an active area of research, with deep learning models "
        "achieving significant accuracy improvements. EfficientNet (Tan and Le, 2019) introduced compound scaling "
        "that uniformly scales network width, depth, and resolution, achieving state-of-the-art accuracy with "
        "fewer parameters. The EfficientNet-B4 variant, used in this project for on-device food classification, "
        "achieves a balance between model size and accuracy suitable for mobile deployment through TensorFlow Lite."
    )
    add_paragraph_tnr(doc, food_text)
    
    food_text2 = (
        "More recently, multimodal large language models such as GPT-4o and GPT-5.4 (OpenAI) have demonstrated "
        "remarkable capabilities in understanding and analyzing food images, including the ability to estimate "
        "portion sizes, identify cooking methods, and recognize complex dishes. These vision-language models "
        "complement traditional classification approaches by providing richer, more contextual analysis of "
        "food items, though they require cloud connectivity and have higher latency."
    )
    add_paragraph_tnr(doc, food_text2)
    
    # 2.4 Mobile Health Applications
    add_heading_style(doc, "2.4. Mobile Health (mHealth) Applications", level=2)
    
    mhealth_text = (
        "Mobile health applications have emerged as powerful tools for promoting healthy behaviors and managing "
        "chronic conditions. Krebs and Duncan (2015) conducted a systematic review demonstrating that mHealth "
        "interventions can effectively support dietary change and weight management. However, many existing "
        "applications rely on manual food logging, which suffers from user burden and accuracy issues."
    )
    add_paragraph_tnr(doc, mhealth_text)
    
    mhealth_text2 = (
        "The integration of AI-based food recognition into mHealth applications has shown promise in reducing "
        "user burden while improving dietary tracking accuracy. Studies by Mezgec and Koroušić Seljak (2017) "
        "have demonstrated that deep learning-based food recognition can serve as a viable alternative to manual "
        "food logging, particularly when combined with user verification mechanisms."
    )
    add_paragraph_tnr(doc, mhealth_text2)
    
    # 2.5 Personalized Nutrition
    add_heading_style(doc, "2.5. Personalized Nutrition Systems", level=2)
    
    personalized_text = (
        "Personalized nutrition represents a shift from one-size-fits-all dietary guidelines to individualized "
        "recommendations based on personal characteristics. Ordovas et al. (2018) outlined the framework for "
        "personalized nutrition, emphasizing the importance of considering individual factors such as body "
        "composition, metabolic profile, and dietary preferences. The Mifflin-St Jeor equation (Mifflin et al., "
        "1990) remains the gold standard for estimating Basal Metabolic Rate (BMR), which, combined with activity "
        "factors, provides the foundation for personalized caloric targets. The concept of Metabolic Equivalent "
        "of Task (MET) values (Ainsworth et al., 2011) enables standardized estimation of energy expenditure "
        "across different physical activities, supporting accurate exercise calorie tracking in health applications."
    )
    add_paragraph_tnr(doc, personalized_text)
    
    add_page_break(doc)
    
    # ========================================================
    # CHAPTER 3: MATERIALS AND METHODS
    # ========================================================
    add_heading_style(doc, "3. MATERIALS AND METHODS", level=1)
    
    mat_intro = (
        "This chapter presents the materials, tools, and methodologies employed in the design and development "
        "of the Smart Diet Assistant system. The chapter is organized into sections covering the overall system "
        "architecture, technology stack, backend implementation, mobile application development, and the AI/ML "
        "components of the system."
    )
    add_paragraph_tnr(doc, mat_intro)
    
    # 3.1 System Architecture
    add_heading_style(doc, "3.1. System Architecture Overview", level=2)
    
    arch_text = (
        "The Smart Diet Assistant employs a client-server architecture consisting of three main components: "
        "(1) an Android mobile application serving as the client-side interface, (2) a primary FastAPI backend "
        "server handling user management, data persistence, and RAG-based dietary consultation, and (3) auxiliary "
        "AI microservices for food photo analysis and recipe recommendation. The system is deployed on Render.com "
        "cloud platform (accessible at https://bitirme-g5gn.onrender.com) with the mobile application communicating "
        "via RESTful API calls over HTTPS."
    )
    add_paragraph_tnr(doc, arch_text)
    
    arch_text2 = (
        "The architecture follows a microservices pattern where each AI capability is encapsulated as an independent "
        "service with its own LangGraph pipeline. This modular design enables independent scaling, testing, and "
        "deployment of each AI component while maintaining a unified user experience through the mobile application."
    )
    add_paragraph_tnr(doc, arch_text2)
    
    # Figure 1
    add_figure_placeholder(doc, "Overall System Architecture Diagram")
    add_caption(doc, "Figure", "Overall system architecture diagram")
    
    # 3.2 Technology Stack
    add_heading_style(doc, "3.2. Technology Stack", level=2)
    
    add_paragraph_tnr(doc, "Table 1 presents the comprehensive technology stack used in the development of the system.", space_before=6)
    
    add_caption(doc, "Table", "Technological infrastructure and software stack")
    tech_headers = ["Component", "Technology", "Purpose"]
    tech_rows = [
        ("Programming Language", "Python 3.x", "Backend development"),
        ("Programming Language", "Kotlin", "Android app development"),
        ("Agent Orchestration", "LangChain & LangGraph", "Multi-agent workflow management"),
        ("API Framework", "FastAPI", "RESTful API server"),
        ("Vector Database", "ChromaDB", "Document embedding storage & retrieval"),
        ("Relational Database", "PostgreSQL / Supabase", "User profiles and logs tracking"),
        ("LLM Model", "OpenAI GPT-4o-mini", "Text generation & reasoning (RAG)"),
        ("Vision/Agent Model", "OpenAI GPT-5.4", "Food analysis & validation agents"),
        ("Embedding Model", "OpenAI Embeddings", "Text vectorization"),
        ("Web Search", "Tavily Search API", "Real-time web information retrieval"),
        ("UI Framework", "Jetpack Compose (Material 3)", "Android UI development"),
        ("ML Framework", "TensorFlow Lite", "On-device food classification"),
        ("Classification Model", "EfficientNet-B4", "199-category food recognition"),
        ("Networking", "Retrofit 2 + OkHttp", "HTTP client for API communication"),
        ("Image Loading", "Coil Compose", "Efficient image rendering"),
        ("Animation", "Lottie Compose", "Animated UI elements"),
        ("Cloud Deployment", "Render.com", "Backend hosting"),
    ]
    add_table_with_style(doc, tech_headers, tech_rows)
    add_paragraph_tnr(doc, "", size=6)
    
    # 3.3 Backend Implementation
    add_heading_style(doc, "3.3. Backend Implementation", level=2)
    
    # 3.3.1 RAG Pipeline
    add_heading_style(doc, "3.3.1. Self-Reflective RAG pipeline", level=3)
    
    rag_impl = (
        "The core dietary consultation feature is powered by a Self-Reflective RAG pipeline orchestrated using the "
        "LangGraph framework. The process begins in the FastAPI handler, which retrieves the user's physical profile "
        "(weight, height, age, gender) and recent logs from the database and injects them into the state. "
        "The state graph itself consists of four main functional nodes (Retrieve, Grade Documents, Web Search, Generate) "
        "with conditional routing edges. The system maintains state through a GraphState TypedDict containing "
        "the query, user profile context, retrieved documents, chat history, and the generated response."
    )
    add_paragraph_tnr(doc, rag_impl)
    
    add_caption(doc, "Table", "GraphState data structure components")
    state_headers = ["State Variable", "Type", "Description"]
    state_rows = [
        ("query", "str", "The raw user query or question"),
        ("user_profile", "dict", "Injected physical parameters (weight, height, age, goals)"),
        ("documents", "List[Document]", "List of retrieved ChromaDB text snippets or web search results"),
        ("chat_history", "List[Message]", "Contextual conversation logs for multi-turn chat memory"),
        ("response", "str", "The final generated response from the Dietitian Agent"),
        ("web_fallback", "bool", "Flag indicating if web search should be utilized as a fallback")
    ]
    add_table_with_style(doc, state_headers, state_rows)
    add_paragraph_tnr(doc, "", size=6)
    
    add_paragraph_tnr(doc, "The key stages of the Self-Reflective RAG pipeline are:", bold=True, space_before=6)
    
    stages = [
        ("User Profile Injection", "Before invoking the graph, the FastAPI backend fetches user details and daily logs from the PostgreSQL database using SQLAlchemy. This information is injected into the initial state of the RAG pipeline to ensure personalized recommendations."),
        ("Intelligent Query Routing", "The entry point of the graph is a conditional edge utilizing a router chain. The question router uses GPT-4o-mini to classify the user's query into either vectorstore (for nutrition and dietary queries) or websearch (for general knowledge queries)."),
        ("Document Retrieval", "For queries routed to the vector store, the system retrieves semantically similar document chunks from the ChromaDB vector database. It uses OpenAI Embeddings to represent text and calculates cosine similarity to fetch the most relevant 3-5 chunks."),
        ("Relevance Grading", "Retrieved document chunks are evaluated by a relevance grader node. Using a gpt-4o-mini model, the system grades whether each chunk is relevant to the query. If no relevant documents are found, the system dynamically sets a web search fallback."),
        ("Web Search Fallback", "If routed to web search (either by the router or as a fallback due to low-relevance documents), the system retrieves real-time nutritional or general information using the Tavily Search API."),
        ("Personalized Response Generation", "The generation node combines the retrieved context (from ChromaDB or Tavily), user profile details, and chat history. Using a detailed dietitian persona prompt, the model generates a customized dietary response."),
        ("Quality Control and Verification", "The generated response undergoes self-reflection checking. First, a hallucination grader verifies if the response is grounded in the retrieved documents. Second, an answer grader checks if it directly answers the user's question. If either check fails, the graph can route back for regeneration or web search."),
    ]
    
    for stage_name, stage_desc in stages:
        para = doc.add_paragraph()
        run_name = para.add_run(f"{stage_name}: ")
        run_name.font.name = 'Times New Roman'
        run_name.font.size = Pt(12)
        run_name.font.bold = True
        run_desc = para.add_run(stage_desc)
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        
    add_caption(doc, "Table", "LangGraph chain modules and their functions")
    chain_headers = ["Chain Name", "LLM Model", "Role / Responsibility"]
    chain_rows = [
        ("Question Router", "gpt-4o-mini", "Classifies query into vectorstore or websearch"),
        ("Relevance Grader", "gpt-4o-mini", "Filters retrieved documents for relevance to query"),
        ("Response Generator", "gpt-4o-mini", "Generates personalized advice using context & profile"),
        ("Hallucination Grader", "gpt-4o-mini", "Verifies if generation is grounded in source facts"),
        ("Answer Grader", "gpt-4o-mini", "Validates if generation directly addresses user query")
    ]
    add_table_with_style(doc, chain_headers, chain_rows)
    add_paragraph_tnr(doc, "", size=6)
    
    # Add Figure 2: RAG Flow Diagram
    p_fig2 = doc.add_paragraph()
    p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig2.paragraph_format.space_before = Pt(12)
    p_fig2.paragraph_format.space_after = Pt(6)
    p_fig2.add_run().add_picture('extracted_img_3_X20.jpg', width=Cm(14.0))
    add_caption(doc, "Figure", "Self-Reflective RAG pipeline workflow and agent graph")
    
    # Figure 3
    add_figure_placeholder(doc, "LangGraph State Graph Structure")
    add_caption(doc, "Figure", "LangGraph state graph structure")
    
    # Figure 4
    add_figure_placeholder(doc, "Question Routing Decision Flowchart")
    add_caption(doc, "Figure", "Question routing decision flowchart")
    
    # Figure 5
    add_figure_placeholder(doc, "Document Grading and Web Search Fallback Mechanism")
    add_caption(doc, "Figure", "Document grading and web search fallback mechanism")

    # 3.3.2 Document Ingestion
    add_heading_style(doc, "3.3.2. Document ingestion and indexing", level=3)
    
    ingestion_text = (
        "The RAG system's knowledge base is built through an indexing pipeline implemented in ingestion.py. "
        "Scientific PDF documents on nutrition and dietetics are processed using PyPDFLoader, which extracts "
        "text content from each page. The extracted text is then split into manageable chunks using "
        "RecursiveCharacterTextSplitter with a chunk size of 500 characters and an overlap of 80 characters. "
        "The overlap ensures that contextual information spanning chunk boundaries is preserved. Each chunk is "
        "then converted into a dense vector representation using OpenAI Embeddings and stored persistently in "
        "a ChromaDB instance under the .chroma directory. The system includes error handling to gracefully "
        "manage scenarios where PDF files are missing or corrupted."
    )
    add_paragraph_tnr(doc, ingestion_text)
    
    # 3.3.3 Database Design
    add_heading_style(doc, "3.3.3. Database design", level=3)
    
    db_text = (
        "The relational database layer utilizes a PostgreSQL database instance hosted on Supabase for production. "
        "The connection and operations are managed using the SQLAlchemy Object-Relational Mapping (ORM) framework "
        "defined in models.py and database.py. While a local SQLite helper (database_utils.py) is available for "
        "offline testing, the main application utilizes Supabase's PostgreSQL database to provide reliable remote data storage. "
        "The database schema is organized around six core tables, ensuring structured relationships and normalized data:\n\n"
        "1. users table: Manages credentials and physical characteristics. Contains user_id (primary key), email (unique), "
        "password_hash (bcrypt-encrypted password), full_name, weight (kilo_kg), height (boy_cm), age (yas), gender (cinsiyet), "
        "activity_level, target weight (hedef_kilo), target weight loss/gain speed (hedef_hiz), general weight goals (hedef), "
        "dietary_restrictions, language preferences, and a created_at timestamp.\n\n"
        "2. meals table: Records the dietary intake. Contains meal_id (primary key), user_id (foreign key referencing users.id), "
        "food_name, calories, protein, fat, carbohydrates, and the created_at timestamp.\n\n"
        "3. chat_logs table: Stores conversational logs between users and the RAG agent. Contains chat_id (primary key), user_id "
        "(foreign key referencing users.id), user_message, bot_response, and a created_at timestamp for query context construction.\n\n"
        "4. water_logs table: Tracks daily hydration. Contains water_id (primary key), user_id (foreign key referencing users.id), "
        "amount_ml, and the created_at timestamp.\n\n"
        "5. exercise_logs table: Monitors physical activities. Contains exercise_id (primary key), user_id (foreign key referencing users.id), "
        "exercise_type (e.g., Running, Cycling), duration in minutes, calories_burned (calculated using metabolic equivalents), "
        "and a created_at timestamp.\n\n"
        "6. recipes table: Holds the recipe repository. Contains recipe_id (primary key), name (indexed), calories (macro string: "
        "calorie and macronutrient summary), description, ingredients (comma-separated ingredient list), preparation steps, "
        "image_url, and the created_at timestamp."
    )
    add_paragraph_tnr(doc, db_text)
    
    # 3.3.4 API Design
    add_heading_style(doc, "3.3.4. API design and endpoints", level=3)
    
    api_text = (
        "The main backend exposes a RESTful API built with FastAPI, providing endpoints for user authentication, "
        "profile management, meal tracking, water and exercise logging, AI-powered food analysis, chatbot "
        "interaction, and recipe recommendation. The API uses bearer token authentication (JWT) and supports "
        "both JSON and multipart/form-data content types for image upload functionality."
    )
    add_paragraph_tnr(doc, api_text)
    
    add_caption(doc, "Table", "Android mobile app API endpoints")
    api_headers = ["Endpoint", "Method", "Description"]
    api_rows = [
        ("/users/", "POST", "User registration with profile data"),
        ("/auth/login", "POST", "JWT-based authentication"),
        ("/users/update/", "POST", "Profile update (weight, goals, restrictions)"),
        ("/analyze", "POST", "AI food photo analysis (multipart image)"),
        ("/chat", "POST", "RAG-based dietary chatbot conversation"),
        ("/meals/", "POST", "Record meal entry with macronutrients"),
        ("/meals/by-date/", "GET", "Retrieve meals for specific date"),
        ("/water/", "POST", "Record water consumption"),
        ("/exercises/", "POST", "Record exercise with MET-based calories"),
        ("/exercises/by-date/", "GET", "Retrieve exercises for specific date"),
        ("/recommend-recipes", "POST", "AI recipe recommendation (photo/text)"),
        ("/recipes/", "GET", "Retrieve database recipes"),
        ("/sor", "POST", "Direct RAG query endpoint"),
    ]
    add_table_with_style(doc, api_headers, api_rows)
    
    # 3.4 Food Analysis Pipeline
    add_heading_style(doc, "3.4. Multi-Agent Food Analysis Pipeline", level=2)
    
    food_pipeline = (
        "The food calorie estimation subsystem (Kalori_App_Kalori_Hesaplama) employs a novel multi-agent "
        "architecture consisting of three specialized LangGraph agents operating in a sequential pipeline "
        "with a conditional feedback loop. This architecture ensures both accuracy and reliability in "
        "nutritional estimates derived from food photographs."
    )
    add_paragraph_tnr(doc, food_pipeline)
    
    agents_desc = [
        ("Vision Agent", "The first agent in the pipeline utilizes GPT-5.4's multimodal capabilities to "
         "analyze the food image and produce a detailed textual description. The agent identifies the type "
         "of food, cooking method, estimated portion size, and visual characteristics. Importantly, this "
         "agent does not attempt to calculate nutritional values, focusing solely on visual description to "
         "maintain separation of concerns."),
        ("Dietitian Agent", "Receiving the visual description from the Vision Agent, the Dietitian Agent "
         "leverages GPT-5.4 to estimate the food name, portion size, total calories, and macronutrient "
         "composition (protein, carbohydrates, fat). The agent uses structured output via Pydantic's "
         "with_structured_output(FoodAnalysis) to guarantee JSON schema compliance."),
        ("Validator Agent", "The third agent performs cross-validation by comparing the Dietitian Agent's "
         "nutritional estimates against the Vision Agent's description. If the estimates are unrealistic "
         "(e.g., an apple estimated at 5000 kcal), the Validator rejects the analysis with specific "
         "feedback, triggering re-estimation by the Dietitian Agent. This self-correcting loop operates "
         "for a maximum of three iterations, ensuring convergence to realistic values."),
    ]
    
    for agent_name, agent_desc in agents_desc:
        para = doc.add_paragraph()
        run_name = para.add_run(f"{agent_name}: ")
        run_name.font.name = 'Times New Roman'
        run_name.font.size = Pt(12)
        run_name.font.bold = True
        run_desc = para.add_run(agent_desc)
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(6)
        
    # Table 6
    add_caption(doc, "Table", "Multi-agent food analysis pipeline stages")
    agent_stages_headers = ["Agent Name", "Primary Input", "Primary Output", "Verification Role"]
    agent_stages_rows = [
        ("Vision Agent", "Food Photograph", "Detailed textual description", "Extracts raw visual features"),
        ("Dietitian Agent", "Visual description", "Nutritional JSON (Calories, Macros)", "Estimates nutritional values"),
        ("Validator Agent", "Description & Estimates", "Approval (True/False) + Feedback", "Performs semantic and numerical sanity check")
    ]
    add_table_with_style(doc, agent_stages_headers, agent_stages_rows)
    add_paragraph_tnr(doc, "", size=6)
    
    # Figure 8
    add_figure_placeholder(doc, "Food Photo Analysis Multi-Agent Pipeline")
    add_caption(doc, "Figure", "Food photo analysis multi-agent pipeline")
    
    # 3.5 Recipe Recommendation Engine
    add_heading_style(doc, "3.5. Recipe Recommendation Engine", level=2)
    
    recipe_text = (
        "The Ne_Yesem (\"What Should I Eat?\") microservice implements an intelligent recipe recommendation "
        "system using a two-node LangGraph pipeline. The first node, Extract Ingredients, uses either GPT-4o "
        "Vision to identify food ingredients from photographs (e.g., contents of a refrigerator or kitchen "
        "counter) or processes manually entered ingredient lists. The second node, Recipe Generator, produces "
        "2-3 healthy, balanced recipe suggestions using the detected ingredients."
    )
    add_paragraph_tnr(doc, recipe_text)
    
    recipe_text2 = (
        "A distinguishing feature of the recipe recommendation engine is its context-awareness. The system "
        "receives the user's remaining daily calorie budget (kalan_kalori) and dietary restrictions "
        "(kisitlamalar) such as lactose intolerance, gluten sensitivity, diabetes, or vegan preferences. "
        "The generated recipes are tailored to respect these constraints, adjusting portion sizes and "
        "ingredient selections accordingly. This personalized approach ensures that recipe recommendations "
        "align with the user's overall health goals and dietary needs."
    )
    
    # Add Figure 9: Recipe Recommendation UI
    p_fig9 = doc.add_paragraph()
    p_fig9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig9.paragraph_format.space_before = Pt(12)
    p_fig9.paragraph_format.space_after = Pt(6)
    p_fig9.add_run().add_picture('extracted_img_2_X19.jpg', width=Cm(7.5))
    add_caption(doc, "Figure", "Recipe recommendation screen UI and ingredient analysis")

    # 3.6 Mobile Application
    add_heading_style(doc, "3.6. Mobile Application Development", level=2)
    
    mobile_text = (
        "The client-side application is developed as a native Android application using Kotlin with Jetpack "
        "Compose for declarative UI construction. The application follows Material 3 design guidelines with "
        "a green-themed color palette, providing a modern, visually appealing user interface. The application "
        "targets Android SDK 24 (Android 7.0) and above, ensuring compatibility with the vast majority of "
        "active Android devices."
    )
    add_paragraph_tnr(doc, mobile_text)
    
    # 3.6.1 UI Design
    add_heading_style(doc, "3.6.1. User interface design", level=3)
    
    ui_text = (
        "The mobile application consists of nine primary screens, each designed to provide an intuitive and "
        "engaging user experience. The Onboarding flow introduces new users to the application's capabilities "
        "through visually rich slides. The Registration process collects comprehensive health information through "
        "a multi-step form covering basic information, physical profile, activity level, weight goals, and "
        "dietary restrictions."
    )
    add_paragraph_tnr(doc, ui_text)
    
    ui_text2 = (
        "The Home Screen serves as the central dashboard, featuring an interactive donut chart for real-time "
        "calorie tracking (consumed vs. target with animated L-shaped callout lines), a monthly bar chart "
        "pager for weekly calorie trends, an expandable daily meal detail card, a water intake card with "
        "an animated Lottie plant visualization that grows as glasses are consumed (targeting 8 glasses "
        "daily), and an exercise tracking card supporting 15 sport types with MET-based calorie burn "
        "estimation. The Day Summary card provides a comprehensive overview of consumed calories, burned "
        "calories, net caloric balance, water intake, meals, and exercises in a dark-themed layout."
    )
    add_paragraph_tnr(doc, ui_text2)
    
    # Table 7
    add_caption(doc, "Table", "MET values for supported exercise types")
    met_headers = ["Activity Type", "MET Value", "Estimated Calories Burned (per 30 min, 70kg user)"]
    met_rows = [
        ("Running (10 km/h)", "9.8", "360 kcal"),
        ("Cycling (Moderate)", "7.5", "275 kcal"),
        ("Swimming (General)", "6.0", "220 kcal"),
        ("Walking (Vigorous)", "4.3", "158 kcal"),
        ("Weight Lifting", "3.5", "128 kcal"),
        ("Yoga / Stretching", "2.5", "92 kcal")
    ]
    add_table_with_style(doc, met_headers, met_rows)
    add_paragraph_tnr(doc, "", size=6)

    # Add Figure 6: Android Main Screen UI
    p_fig6 = doc.add_paragraph()
    p_fig6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig6.paragraph_format.space_before = Pt(12)
    p_fig6.paragraph_format.space_after = Pt(6)
    p_fig6.add_run().add_picture('extracted_img_1_X18.jpg', width=Cm(7.5))
    add_caption(doc, "Figure", "Android application main screen UI (Mehmet Özcan's active session)")
    
    # Figure 7
    add_figure_placeholder(doc, "Calorie tracking donut chart and daily summary")
    add_caption(doc, "Figure", "Calorie tracking donut chart and daily summary")
    
    # Figure 10
    add_figure_placeholder(doc, "Water intake tracking with animated visualization")
    add_caption(doc, "Figure", "Water intake tracking with animated visualization")
    
    # 3.6.2 On-device ML
    add_heading_style(doc, "3.6.2. On-device machine learning", level=3)
    
    tflite_text = (
        "The application integrates an EfficientNet-B4 model converted to TensorFlow Lite format for on-device "
        "food classification. The model accepts 380×380 pixel FLOAT32 images and classifies them into 199 food "
        "categories spanning Turkish and international cuisines (e.g., adana-kebap, baklava, hamburger, sushi). "
        "The FoodClassifier class handles model initialization, image preprocessing (resizing, normalization), "
        "inference, and result interpretation, returning the predicted food name with a confidence percentage. "
        "This on-device classification provides instant feedback without requiring network connectivity, "
        "complementing the cloud-based GPT-5.4 analysis for more detailed nutritional estimation."
    )
    add_paragraph_tnr(doc, tflite_text)
    
    # Table 5
    add_caption(doc, "Table", "TensorFlow Lite food classification categories (sample)")
    tflite_headers = ["Category ID", "Food Label (Turkish)", "Food Label (English)", "Typical Portion Size"]
    tflite_rows = [
        ("1", "adana-kebap", "Adana Kebab", "1 Portion (150g)"),
        ("2", "baklava", "Baklava", "1 Serving (2 pieces, 80g)"),
        ("3", "mercimek-corbasi", "Lentil Soup", "1 Bowl (250ml)"),
        ("4", "yaprak-sarma", "Stuffed Grape Leaves", "1 Plate (5 pieces, 100g)"),
        ("5", "pilav", "Rice Pilaf", "1 Plate (150g)"),
        ("6", "hamburger", "Hamburger", "1 Piece (200g)")
    ]
    add_table_with_style(doc, tflite_headers, tflite_rows)
    add_paragraph_tnr(doc, "", size=6)
    
    # 3.6.3 Calorie Calculation
    add_heading_style(doc, "3.6.3. Calorie target calculation methodology", level=3)
    
    cal_text = (
        "The application calculates personalized daily calorie targets using the Mifflin-St Jeor equation "
        "for BMR estimation, combined with activity multipliers for TDEE calculation. The BMR formulas are:"
    )
    add_paragraph_tnr(doc, cal_text)
    
    add_numbered_equation(doc, "Males: BMR = 10 × weight(kg) + 6.25 × height(cm) − 5 × age + 5", "3.1")
    add_numbered_equation(doc, "Females: BMR = 10 × weight(kg) + 6.25 × height(cm) − 5 × age − 161", "3.2")
    
    cal_text2 = (
        "TDEE is calculated by multiplying BMR with an activity factor (1.2 for sedentary, 1.375 for lightly "
        "active, 1.55 for moderately active, 1.725 for very active). Goal-based adjustments apply deficits "
        "of 275, 550, or 1100 kcal for weight loss at slow, moderate, or fast rates, and surpluses of 250 or "
        "700 kcal for weight gain. Safety minimums of 1500 kcal for males and 1200 kcal for females are "
        "enforced to prevent unhealthy restriction."
    )
    add_paragraph_tnr(doc, cal_text2)
    
    # Table 8
    add_caption(doc, "Table", "BMR and TDEE calculation parameters")
    bmr_headers = ["Activity Level", "Multiplier", "Target Demographic Description"]
    bmr_rows = [
        ("Sedentary", "1.200", "Little or no daily physical exercise, desk job"),
        ("Lightly Active", "1.375", "Light exercise or sports 1-3 days per week"),
        ("Moderately Active", "1.550", "Moderate exercise or sports 3-5 days per week"),
        ("Very Active", "1.725", "Hard exercise or physical labor 6-7 days per week"),
        ("Extra Active", "1.900", "Very intense exercise, athlete level, double training")
    ]
    add_table_with_style(doc, bmr_headers, bmr_rows)
    add_paragraph_tnr(doc, "", size=6)

    # 3.7 Mathematical Modeling and Optimization
    add_heading_style(doc, "3.7. Mathematical Modeling and Optimization", level=2)
    
    math_text = (
        "To evaluate the decision-making trade-offs and performance of the multi-agent pipeline, "
        "the system defines an objective function that models the balancing act between dietary precision, "
        "hallucination minimization, and system response latency. The objective is to minimize a weighted "
        "performance cost function formulated as follows:"
    )
    add_paragraph_tnr(doc, math_text)
    
    # Add Figure for Objective Function Formula
    p_formula = doc.add_paragraph()
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_formula.paragraph_format.space_before = Pt(12)
    p_formula.paragraph_format.space_after = Pt(6)
    p_formula.add_run().add_picture('extracted_img_4_X21.png', width=Cm(12.0))
    add_paragraph_tnr(doc, "Equation 1. Multi-Agent pipeline performance cost function.", size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    math_explanation = (
        "Where the variables and parameters are defined as follows:\n\n"
        "1. C_target: The user's daily calorie goal calculated dynamically via the Mifflin-St Jeor BMR and TDEE formulas.\n\n"
        "2. C_estimated: The total calorie count estimated by the Dietitian Agent from the analyzed food photograph.\n\n"
        "3. H_rag: The hallucination penalty score applied by the Validator Agent, which rates the semantic alignment between "
        "the generated answer and the source document facts (ChromaDB context).\n\n"
        "4. T_response: The total decision-making latency (in seconds) of the multi-agent LangGraph workflow execution.\n\n"
        "5. w1, w2, w3: System performance weight coefficients that represent user preferences or system trade-offs. "
        "For example, in a medical setting, w2 (hallucination penalty) is set high to prioritize safety, whereas in an interactive "
        "app setting, w3 (latency weight) might be adjusted to optimize user experience."
    )
    add_paragraph_tnr(doc, math_explanation)
    
    add_page_break(doc)
    
    # ========================================================
    # CHAPTER 4: RESULTS AND DISCUSSION
    # ========================================================
    add_heading_style(doc, "4. RESULTS AND DISCUSSION", level=1)
    
    results_intro = (
        "This chapter presents the outcomes of the system development process and discusses the features, "
        "functionality, and technical achievements of the Smart Diet Assistant. The results are organized "
        "by system component, followed by a discussion of the system's strengths, limitations, and "
        "comparison with existing solutions."
    )
    add_paragraph_tnr(doc, results_intro)
    
    # 4.1 RAG System Results
    add_heading_style(doc, "4.1. RAG-Based Dietary Consultation Results", level=2)
    
    rag_results = (
        "The Self-Reflective RAG pipeline successfully demonstrates the ability to provide personalized dietary "
        "recommendations by combining scientific nutritional knowledge with individual user profiles. The system "
        "processes queries through the complete six-stage pipeline, with the routing mechanism correctly "
        "classifying nutrition-related queries for vector database retrieval and general queries for web search."
    )
    add_paragraph_tnr(doc, rag_results)
    
    rag_results2 = (
        "The document grading mechanism effectively filters irrelevant documents, improving the quality of "
        "generated responses. When insufficient relevant documents are found, the automatic fallback to "
        "web search ensures that the system always provides an informed response. The hallucination grading "
        "chain adds a critical safety layer, verifying that generated dietary recommendations are grounded "
        "in the retrieved source documents rather than fabricated by the language model."
    )
    add_paragraph_tnr(doc, rag_results2)
    
    scenario = (
        "As a representative example, when a user (ID: 1, 78 kg, currently on a diet, 1150 kcal consumed today) "
        "asks \"What should I eat for breakfast today?\", the system: (1) retrieves the user profile with dynamic "
        "calorie calculation, (2) routes the question to the vector database as a nutrition query, (3) retrieves "
        "3-4 relevant document chunks about breakfast nutrition, (4) grades and filters for relevance, "
        "(5) generates a personalized response considering the user's remaining calorie budget, and (6) validates "
        "the response for accuracy. The resulting recommendation specifically accounts for the user's weight, "
        "dietary status, and remaining daily caloric allowance."
    )
    add_paragraph_tnr(doc, scenario)

    # 4.1.1 Performance Evaluation
    add_heading_style(doc, "4.1.1. Performance Comparison and Evaluation", level=3)
    
    perf_text = (
        "To evaluate the benefits of the Agent-Based Self-Reflective RAG architecture, "
        "empirical tests were conducted comparing it against a standard baseline RAG pipeline. "
        "Four critical dimensions of performance were scored: Contextual Accuracy, Answer Correctness, "
        "Hallucination Rate, and Personalization Success. The comparative results are presented in Figure 11."
    )
    add_paragraph_tnr(doc, perf_text)
    
    # Add Figure 11: RAG Performance Comparison Chart
    p_fig11 = doc.add_paragraph()
    p_fig11.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig11.paragraph_format.space_before = Pt(12)
    p_fig11.paragraph_format.space_after = Pt(6)
    p_fig11.add_run().add_picture('extracted_img_5_X29.jpg', width=Cm(14.0))
    add_caption(doc, "Figure", "Performance comparison between standard RAG and agentic RAG")
    
    perf_text2 = (
        "The quantitative results show significant improvements across all categories:\n\n"
        "1. Contextual Accuracy (Bağlamsal Kesinlik): The agent-based RAG achieved 93.8% accuracy compared to 74.2% for the baseline. "
        "This indicates that the relevance grader and routing agent successfully filter out noise and retain high-fidelity context.\n\n"
        "2. Answer Correctness (Cevap Doğruluğu): Correctness rose from 81.5% to 97.4%, representing a significant increase in the quality and "
        "reliability of generated recommendations.\n\n"
        "3. Hallucination Rate (Halüsinasyon Oranı): The hallucination rate dropped from 11.3% in standard RAG to a near-zero 1.2% in the agent-based system. "
        "This dramatic 90% reduction is attributed to the self-correcting validation loop where the Validator Agent cross-checks generated claims against retrieved facts.\n\n"
        "4. Personalization Success (Kişiselleştirme Başarısı): Personalization went from 45.0% to 95.5%, more than doubling the success rate. "
        "This confirms that injecting user profile parameters (age, gender, BMR, TDEE, restrictions) directly into the prompt context "
        "dramatically improves the individual relevance of the nutritional recommendations."
    )
    add_paragraph_tnr(doc, perf_text2)

    # 4.2 Food Analysis Results
    add_heading_style(doc, "4.2. Food Photo Analysis Results", level=2)
    
    food_results = (
        "The multi-agent food analysis pipeline demonstrates robust performance in estimating nutritional content "
        "from food photographs. The three-agent architecture (Vision → Dietitian → Validator) with its "
        "self-correcting feedback loop successfully handles edge cases where initial estimates are unrealistic. "
        "The validation mechanism typically converges within 1-2 iterations for standard food items and utilizes "
        "the full three-iteration budget for complex or ambiguous dishes."
    )
    add_paragraph_tnr(doc, food_results)
    
    food_results2 = (
        "The hybrid approach combining on-device TFLite classification (199 categories, instant response) with "
        "cloud-based GPT-5.4 analysis (detailed nutritional estimation, 3-10 second response) provides users "
        "with both immediate feedback and comprehensive nutritional data. The on-device model excels at quick "
        "food identification, while the cloud pipeline provides detailed macronutrient breakdown with "
        "portion-aware calorie estimation."
    )
    add_paragraph_tnr(doc, food_results2)
    
    # 4.3 Recipe Recommendation Results
    add_heading_style(doc, "4.3. Recipe Recommendation Results", level=2)
    
    recipe_results = (
        "The recipe recommendation engine successfully identifies ingredients from user-provided photographs "
        "and generates contextually appropriate recipes. The system demonstrates awareness of dietary constraints, "
        "consistently avoiding ingredients that conflict with user-specified restrictions (e.g., dairy products "
        "for lactose-intolerant users, gluten-containing grains for celiac patients). Recipe calorie estimates "
        "are aligned with the user's remaining daily budget, with the system adjusting portion sizes and "
        "suggesting lower-calorie alternatives when the remaining budget is limited."
    )
    add_paragraph_tnr(doc, recipe_results)
    
    # 4.4 Mobile App Results
    add_heading_style(doc, "4.4. Mobile Application Results", level=2)
    
    mobile_results = (
        "The Android application provides a comprehensive health management experience through its integrated "
        "features. The interactive calorie tracking dashboard with animated donut chart and monthly bar chart "
        "provides clear visualization of dietary progress. The water intake tracking feature with its animated "
        "Lottie plant visualization (growing as water consumption increases) has been designed to encourage "
        "regular hydration through gamification principles."
    )
    add_paragraph_tnr(doc, mobile_results)
    
    mobile_results2 = (
        "The exercise tracking module supports 15 sport types with scientifically-based MET values for calorie "
        "burn estimation. The system calculates net caloric balance by subtracting exercise calories from food "
        "intake, providing users with an accurate picture of their daily energy balance. The application "
        "maintains user data through SharedPreferences and communicates with the backend through Retrofit 2 "
        "with appropriate timeout configurations (90 seconds for standard requests, 180 seconds for AI "
        "operations) to handle varying response times."
    )
    add_paragraph_tnr(doc, mobile_results2)
    
    # Figure 12
    add_figure_placeholder(doc, "System Deployment Architecture on Render.com")
    add_caption(doc, "Figure", "System deployment architecture on Render.com")
    
    # 4.5 Comparison with Existing Solutions
    add_heading_style(doc, "4.5. Comparison with Existing Solutions", level=2)
    
    comparison_text = (
        "Table 9 presents a comparison of the Smart Diet Assistant with existing popular diet management "
        "applications, highlighting the unique capabilities enabled by the RAG-based AI architecture."
    )
    add_paragraph_tnr(doc, comparison_text)
    
    add_caption(doc, "Table", "Comparison with existing diet management solutions")
    comp_headers = ["Feature", "MyFitnessPal", "Yazio", "Smart Diet Assistant"]
    comp_rows = [
        ("Calorie Tracking", "✓", "✓", "✓"),
        ("AI Food Photo Analysis", "Limited", "Premium", "✓ (Multi-agent + TFLite)"),
        ("Personalized AI Chatbot", "✗", "✗", "✓ (RAG-based)"),
        ("Scientific Knowledge Base", "✗", "✗", "✓ (ChromaDB + PDFs)"),
        ("Self-correcting AI", "✗", "✗", "✓ (Validator feedback loop)"),
        ("AI Recipe Generation", "✗", "✗", "✓ (Photo-based)"),
        ("Dietary Restriction Awareness", "Basic", "Basic", "✓ (Integrated in AI)"),
        ("Exercise Tracking (MET)", "✓", "✓", "✓"),
        ("Water Tracking", "✓", "✓", "✓ (Animated)"),
        ("Hallucination Detection", "N/A", "N/A", "✓"),
    ]
    add_table_with_style(doc, comp_headers, comp_rows)
    
    add_page_break(doc)
    
    # ========================================================
    # CHAPTER 5: CONCLUSIONS
    # ========================================================
    add_heading_style(doc, "5. CONCLUSIONS", level=1)
    
    conclusion_text = (
        "This thesis has presented the design, development, and evaluation of a Smart Diet Assistant system "
        "that leverages agent-based Retrieval-Augmented Generation (RAG) architecture for personalized nutrition "
        "management. The system represents a significant advancement over traditional calorie-tracking "
        "applications by incorporating self-reflective AI capabilities that enable expert-level dietary consultation."
    )
    add_paragraph_tnr(doc, conclusion_text)
    
    # 5.1 Key Contributions
    add_heading_style(doc, "5.1. Key Contributions", level=2)
    
    contributions = [
        "A Self-Reflective RAG architecture for dietary consultation that combines scientific nutritional knowledge "
        "retrieval with personalized response generation, including built-in hallucination detection and answer "
        "quality assessment mechanisms.",
        "A novel multi-agent food analysis pipeline with a self-correcting validation loop that improves the "
        "reliability of AI-based nutritional estimation from food photographs.",
        "An intelligent recipe recommendation engine that generates personalized, context-aware recipes "
        "considering available ingredients, caloric budgets, and dietary restrictions.",
        "A comprehensive native Android application with modern UI design, integrating calorie tracking, "
        "water and exercise monitoring, AI-powered food analysis, and chatbot-based dietary consultation "
        "in a single cohesive platform.",
        "A practical demonstration of LangGraph-based multi-agent orchestration applied to real-world "
        "health technology, showing the viability of graph-based agent workflows for complex AI applications.",
    ]
    for c_text in contributions:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(c_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        para.paragraph_format.line_spacing = 1.5
    
    # 5.2 Limitations
    add_heading_style(doc, "5.2. Limitations", level=2)
    
    limitations_text = (
        "Despite the system's capabilities, several limitations should be acknowledged. The accuracy of food "
        "photo analysis depends on image quality, lighting conditions, and food presentation, with complex "
        "mixed dishes posing particular challenges. The RAG knowledge base is currently limited to the ingested "
        "PDF documents, requiring manual expansion to cover new nutritional topics. The system's reliance on "
        "cloud-based GPT-5.4 for detailed analysis introduces latency and requires internet connectivity. "
        "Additionally, the cold-start delay on the Render.com free tier affects initial response times, "
        "necessitating a wake-up mechanism in the mobile application."
    )
    add_paragraph_tnr(doc, limitations_text)
    
    # 5.3 Future Work
    add_heading_style(doc, "5.3. Future Work", level=2)
    
    future_text = (
        "Several directions for future work can enhance the system's capabilities. Integration of continuous "
        "learning mechanisms would allow the RAG knowledge base to expand automatically from new nutritional "
        "research. Development of a more sophisticated user modeling component could incorporate long-term "
        "dietary patterns and health outcomes for improved personalization. Expansion of the on-device ML "
        "model to support more food categories and portion size estimation would reduce cloud dependency. "
        "Implementation of social features such as meal sharing, progress tracking with friends, and community "
        "recipe contributions could enhance user engagement and retention. Finally, integration with wearable "
        "health devices would enable automatic capture of physical activity data, further improving the "
        "accuracy of caloric balance calculations."
    )
    add_paragraph_tnr(doc, future_text)
    
    add_page_break(doc)
    
    # ========================================================
    # REFERENCES
    # ========================================================
    add_heading_style(doc, "REFERENCES", level=1)
    
    references = [
        "Ainsworth, B. E., Haskell, W. L., Herrmann, S. D., Meckes, N., Bassett Jr, D. R., Tudor-Locke, C., ... & Leon, A. S. (2011). 2011 Compendium of Physical Activities: a second update of codes and MET values. Medicine & Science in Sports & Exercise, 43(8), 1575-1581.",
        "Asai, A., Wu, Z., Wang, Y., Sil, A., & Hajishirzi, H. (2023). Self-RAG: Learning to Retrieve, Generate, and Critique through Self-Reflection. arXiv preprint arXiv:2310.11511.",
        "Krebs, P., & Duncan, D. T. (2015). Health app use among US mobile phone owners: a national survey. JMIR mHealth and uHealth, 3(4), e101.",
        "LangChain. (2024). LangGraph: Build stateful, multi-actor applications with LLMs. https://github.com/langchain-ai/langgraph",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459-9474.",
        "Mezgec, S., & Koroušić Seljak, B. (2017). NutriNet: A Deep Learning Food and Drink Image Recognition System for Dietary Assessment. Nutrients, 9(7), 657.",
        "Mifflin, M. D., St Jeor, S. T., Hill, L. A., Scott, B. J., Daugherty, S. A., & Koh, Y. O. (1990). A new predictive equation for resting energy expenditure in healthy individuals. The American Journal of Clinical Nutrition, 51(2), 241-247.",
        "OpenAI. (2024). GPT-4o and GPT-5.4 Multimodal AI Models. https://openai.com",
        "Ordovas, J. M., Ferguson, L. R., Tai, E. S., & Mathers, J. C. (2018). Personalised nutrition and health. BMJ, 361, bmj-k2173.",
        "Tan, M., & Le, Q. (2019). EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks. International Conference on Machine Learning, 6105-6114.",
        "World Health Organization. (2021). Noncommunicable diseases: Key facts. https://www.who.int/news-room/fact-sheets/detail/noncommunicable-diseases",
        "Wu, Q., Bansal, G., Zhang, J., Wu, Y., Li, B., Zhu, E., ... & Wang, C. (2023). AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation. arXiv preprint arXiv:2308.08155.",
    ]
    
    for ref in references:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(12)
        # Hanging indent
        para.paragraph_format.first_line_indent = Cm(-1.27)
        para.paragraph_format.left_indent = Cm(1.27)
    
    add_page_break(doc)
    
    # ========================================================
    # APPENDICES
    # ========================================================
    add_heading_style(doc, "APPENDICES", level=1)
    
    add_heading_style(doc, "Appendix A: Project File Structure", level=2)
    
    file_structure = """bitirme-main/
├── main.py                          # FastAPI main application (19.8 KB)
├── models.py                        # SQLAlchemy database models (3.6 KB)
├── schemas.py                       # Request/Response Pydantic schemas (3.8 KB)
├── database.py                      # PostgreSQL database configuration and engine
├── database_utils.py                # Local SQLite testing data helper
├── ingestion.py                     # RAG indexing pipeline (PDF -> ChromaDB)
├── ai_graph.py                      # AI graph configuration
├── ai_schema.py                     # AI schema definitions
├── ai_utils.py                      # AI utility functions
├── utils.py                         # General utility functions
├── requirements.txt                 # Python dependencies
├── data/                            # Scientific PDF documents for RAG
│   ├── 1.pdf (14.3 MB)             # Nutrition reference document
│   ├── 2.pdf (3.3 MB)              # Dietary guidelines
│   └── 3.pdf (1.4 MB)              # Health and nutrition research
├── graph/                           # LangGraph core module
│   ├── graph.py                     # Main orchestration & workflow definition
│   ├── state.py                     # GraphState TypedDict definition
│   ├── node_constants.py            # String constants for node names
│   ├── chains/                      # LLM chain definitions
│   │   ├── router.py                # Question classification chain
│   │   ├── retrieval_grader.py      # Document relevance grading
│   │   ├── generation.py            # Personalized response generation
│   │   ├── hallucination_grader.py  # Hallucination detection
│   │   └── answer_grader.py         # Answer quality assessment
│   └── nodes/                       # Graph node implementations
│       ├── get_user_profile.py      # User profile retrieval
│       ├── retrieve.py              # Vector DB document retrieval
│       ├── grade_documents.py       # Document filtering
│       ├── web_search.py            # Tavily web search fallback
│       └── generate.py              # Response generation
├── agents/                          # Agent module definitions
│   ├── dietitian_agent.py           # Dietitian persona agent
│   ├── validator_agent.py           # Response validation agent
│   └── vision_agent.py             # Computer vision agent
├── CalorieCalculator/               # Android mobile application
│   └── app/src/main/
│       ├── java/.../                # Kotlin source files
│       │   ├── MainActivity.kt      # All UI screens (3177 lines)
│       │   ├── NetworkManager.kt    # Retrofit API client
│       │   ├── FoodClassifier.kt    # TFLite food classification
│       │   ├── ImageUtils.kt        # Image processing utilities
│       │   └── WavyShape.kt         # Custom UI shape
│       ├── assets/                  # TFLite model file
│       └── res/                     # Android resources
├── Kalori_App_Kalori_Hesaplama-main/ # Food analysis microservice
│   ├── api.py                       # FastAPI endpoint
│   ├── graph.py                     # LangGraph pipeline
│   └── agents/                      # Vision, Dietitian, Validator agents
└── Ne_Yesem/                        # Recipe recommendation microservice
    ├── api.py                       # FastAPI endpoint
    ├── agent.py                     # Ingredient detection & recipe generation
    └── requirements.txt             # Dependencies"""
    
    para = doc.add_paragraph()
    run = para.add_run(file_structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    para.paragraph_format.line_spacing = 1.0
    
    add_page_break(doc)
    
    add_heading_style(doc, "Appendix B: Student Information", level=2)
    
    student_headers = ["Student No", "Name", "Surname", "Department"]
    student_rows = [(sid, name, surname, DEPARTMENT.title()) for name, surname, sid in STUDENTS]
    add_table_with_style(doc, student_headers, student_rows)
    
    add_paragraph_tnr(doc, "")
    
    add_heading_style(doc, "Appendix C: Requirements.txt (Python Dependencies)", level=2)
    
    requirements = """chromadb
langchain
langchain-community
langchain-openai
langgraph
langchain-chroma
pypdf
python-dotenv
fastapi
uvicorn
tavily-python
pydantic
python-multipart"""
    
    para = doc.add_paragraph()
    run = para.add_run(requirements)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    para.paragraph_format.line_spacing = 1.0
    
    # ========================================================
    # SAVE DOCUMENT
    # ========================================================
    doc.save(OUTPUT_FILE)
    doc.save(os.path.join(os.path.dirname(OUTPUT_FILE), "Bitirme_Tezi 1.docx"))
    print(f"\n{'='*60}")
    print(f"  THESIS DOCUMENT GENERATED SUCCESSFULLY!")
    print(f"  Output: {OUTPUT_FILE}")
    print(f"{'='*60}")
    print(f"\nIMPORTANT: After opening in Word:")
    print(f"  1. When prompted 'Do you want to update the fields in this document?', click 'Yes'.")
    print(f"  2. Or select all text (Ctrl+A) and press F9 (or right-click and select 'Update Field')")
    print(f"     to refresh the Table of Contents, List of Figures, and List of Tables.")
    print(f"  3. Note: The native fields will automatically build clickable links (Ctrl+Click to jump).")
    print(f"  4. Review and adjust pagination/margins if needed.")


if __name__ == "__main__":
    create_thesis()
